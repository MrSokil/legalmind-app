import streamlit as st
import os
import requests
import time
import sqlite3
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from pypdf import PdfReader
from io import BytesIO

load_dotenv()

# --- НАЛАШТУВАННЯ СТОРІНКИ ---
st.set_page_config(page_title="LegalMind AI", page_icon="⚖️", layout="centered")

# --- CUSTOM CSS ДЛЯ ТЕМНОЇ ТЕМИ ТА ДИЗАЙНУ ---
st.markdown("""
    <style>
    /* Базові налаштування фону */
    .stApp {
        background-color: #0e1117;
    }
    
    /* Заголовки */
    h1, h2, h3 {
        color: #58a6ff !important;
        font-family: 'Inter', sans-serif;
    }
    
    /* Картки результатів (адаптовано під темну тему) */
    .result-card {
        background-color: #161b22;
        padding: 25px;
        border-radius: 12px;
        border: 1px solid #30363d;
        color: #c9d1d9;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
        margin-top: 20px;
        line-height: 1.6;
    }
    
    /* Кнопки в бічній панелі */
    .stSidebar [data-testid="stBaseButton-secondary"] {
        border-radius: 8px;
        border: 1px solid #30363d;
        background-color: #21262d;
        color: #c9d1d9;
        transition: all 0.2s;
        text-align: left;
    }
    
    .stSidebar [data-testid="stBaseButton-secondary"]:hover {
        border-color: #58a6ff;
        color: #58a6ff;
    }
    
    /* Головна кнопка аналізу */
    div.stButton > button:first-child {
        background: linear-gradient(135deg, #1f6feb 0%, #388bfd 100%);
        color: white;
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: 600;
        border: none;
    }
    
    /* Повідомлення про помилки та інфо */
    .stAlert {
        background-color: #161b22;
        border: 1px solid #30363d;
    }
    </style>
    """, unsafe_allow_html=True)

# Отримання ключа
api_key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY")
client = Groq(api_key=api_key)

# --- РОБОТА З БАЗОЮ ДАНИХ ---
def init_db():
    conn = sqlite3.connect('legalmind_history.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS history 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  title TEXT, result TEXT, content TEXT, query TEXT, timestamp DATETIME DEFAULT CURRENT_TIMESTAMP)''')
    conn.commit()
    conn.close()

def save_to_history(title, result, content, query):
    conn = sqlite3.connect('legalmind_history.db')
    c = conn.cursor()
    c.execute("INSERT INTO history (title, result, content, query) VALUES (?, ?, ?, ?)", 
              (title, result, content, query))
    conn.commit()
    conn.close()

def clear_db():
    if os.path.exists('legalmind_history.db'):
        os.remove('legalmind_history.db')
    init_db()

init_db()

# --- СТАН СЕСІЇ ---
if 'analysis_result' not in st.session_state: st.session_state.analysis_result = None
if 'full_content' not in st.session_state: st.session_state.full_content = ""
if 'search_query' not in st.session_state: st.session_state.search_query = ""

# --- ОБРОБКА ФАЙЛІВ ТА URL ---
def read_file(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext == '.txt': return uploaded_file.read().decode("utf-8")
    elif ext == '.docx':
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext == '.pdf':
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages])
    return None

def read_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, timeout=15, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        for s in soup(["script", "style", "header", "footer", "nav"]): s.decompose()
        return soup.get_text(separator=' ', strip=True)
    except Exception as e:
        st.error(f"Помилка посилання: {e}")
        return None

def create_docx(analysis_text):
    doc = Document()
    doc.add_heading('ЮРИДИЧНИЙ ЗВІТ LEGALMIND AI', 0)
    doc.add_paragraph(analysis_text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def auto_determine_mode(text_preview):
    try:
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "Ти — юридичний аналітик. Визнач тип документа. Відповідай ТІЛЬКИ ОДНИМ СЛОВОМ: 'Суд', 'Договір', 'Корпоратив' або 'Загальне'."},
                {"role": "user", "content": text_preview[:3000]}
            ],
            model="llama-3.1-8b-instant",
        )
        detected = response.choices[0].message.content.strip().replace('.', '').replace('"', '').replace("'", "")
        mapping = {
            "Суд": "Судова практика та Позови",
            "Договір": "Аналіз договору",
            "Корпоратив": "Корпоративні документи",
            "Загальне": "Загальна консультація"
        }
        return mapping.get(detected, "Загальна консультація")
    except:
        return "Загальна консультація"

def analyze_long_text(full_text, mode_prompt):
    chunk_size = 7500 
    chunks = [full_text[i:i + chunk_size] for i in range(0, len(full_text), chunk_size)]
    summaries = []
    p_bar = st.progress(0)
    for idx, chunk in enumerate(chunks):
        p_bar.progress((idx + 1) / len(chunks))
        try:
            response = client.chat.completions.create(
                messages=[{"role": "system", "content": "Ти помічник юриста. Випиши ключові юридичні факти."}, {"role": "user", "content": chunk}],
                model="llama-3.1-8b-instant",
            )
            summaries.append(response.choices[0].message.content)
            if idx < len(chunks) - 1: time.sleep(2.1)
        except Exception as e:
            if "rate_limit" in str(e): time.sleep(10)
            else: break
    
    combined = "\n".join(summaries)
    final = client.chat.completions.create(
        messages=[{"role": "system", "content": mode_prompt}, {"role": "user", "content": combined[:18000]}],
        model="llama-3.1-8b-instant",
    )
    p_bar.empty()
    return final.choices[0].message.content

# --- ПРОМПТИ (ВИПРАВЛЕНО ТА ДОПОВНЕНО) ---
prompts = {
    "Судова практика та Позови": """Ти — провідний адвокат із 20-річним стажем. Твоє завдання: знайти слабкі місця в матеріалах та підсилити позицію клієнта. 
    
    Сформуй звіт за наступною суворою структурою:

    ### 🛡️ СТРАТЕГІЧНИЙ АНАЛІЗ ПОЗИЦІЇ
    - **Сильні сторони**: Які факти грають на користь клієнта?
    - **Критичні ризики**: Де позиція найбільш вразлива?

    ### ⚖️ ПЕРЕВІРКА ПРОЦЕСУАЛЬНОЇ ЧИСТОТИ (КЛЮЧОВЕ)
    - Проаналізуй, чи не було порушено процедуру (наприклад, для ст. 130 КУпАП — чи були свідки, чи сертифікований Драгер, чи дотримано строки).
    - Вкажи конкретні статті (наприклад, ст. 251, 268 КУпАП), які допоможуть визнати докази недопустимими.

    ### 📖 РЕКОМЕНДОВАНА ЮРИДИЧНА БАЗА
    - Склади список статей законів та постанов Пленуму ВС, які ТРЕБА додати до позову.
    - Для кожної статті напиши 1 речення: "Це дозволить нам довести, що..."

    ### 📝 ПЛАН ДІЙ (STEP-BY-STEP)
    - Які конкретно клопотання подати (про виклик свідків, про витребування відео з боді-камер тощо)?
    - Яких документів не вистачає в матеріалах справи прямо зараз?

    **Стиль відповіді:** Професійний, гострий, орієнтований на результат.""",

    "Аналіз договору": """Ти — юридичний аудитор. Твоя мета — знайти "пастки" в договорі.
    ### 🚩 ЧЕРВОНІ ПРАПОРЦІ (РИЗИКИ)
    - Знайди приховані штрафи, пеню та умови одностороннього розірвання.
    ### ⚖️ ПОСИЛЕННЯ ПОЗИЦІЇ
    - Які статті ЦКУ/ГКУ захистять клієнта у разі форс-мажору (наприклад, ст. 617 ЦКУ)?
    ### ✏️ РЕКОМЕНДОВАНІ ПРАВКИ
    - Напиши готові формулювання пунктів, які варто змінити (Було -> Стало).""",

    "Корпоративні документи": """Ти корпоративний юрист. Проаналізуй документ на відповідність Статуту та законодавству про господарські товариства. Вияви ризики для засновників та директора.""",

    "Загальна консультація": """Ти професійний юрист. Надай розгорнуту консультацію з посиланням на норми права (Кодекси, Закони). Відповідай структуровано та чітко."""
}

# --- SIDEBAR ---
with st.sidebar:
    st.markdown("### 🏛️ Управління")
    if st.button("🗑️ Очистити історію", use_container_width=True):
        clear_db()
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.rerun()
    
    st.markdown("---")
    st.subheader("📜 Історія аналізів")
    
    conn = sqlite3.connect('legalmind_history.db')
    c = conn.cursor()
    c.execute("SELECT id, title, result, content, query FROM history ORDER BY id DESC LIMIT 15")
    rows = c.fetchall()
    conn.close()

    for h_id, h_title, h_res, h_cont, h_qry in rows:
        if st.button(f"📄 {h_title[:22]}...", key=f"h_{h_id}", use_container_width=True):
            st.session_state.analysis_result = h_res
            st.session_state.full_content = h_cont
            st.session_state.search_query = h_qry
            st.rerun()

# --- MAIN UI ---
st.title("⚖️ LegalMind AI")
st.caption("Система інтелектуального юридичного аналізу та стратегічного планування")

tabs = st.tabs(["📥 Вхідні дані", "📋 Висновки аналізу"])

with tabs[0]:
    source = st.radio("Оберіть метод завантаження:", ["📎 Файл (PDF/DOCX)", "🔗 Посилання (URL)"], horizontal=True)
    content = ""

    if "Файл" in source:
        uploaded_file = st.file_uploader("Завантажте документ", type=['pdf', 'docx', 'txt'])
        if uploaded_file: content = read_file(uploaded_file)
    else:
        url_input = st.text_input("Вкажіть посилання на документ або рішення:")
        if url_input: content = read_url(url_input)

    if content:
        if st.button("🚀 ЗАПУСТИТИ АНАЛІЗ", use_container_width=True):
            try:
                with st.spinner('⚖️ AI формує юридичну стратегію...'):
                    mode = auto_determine_mode(content)
                    st.toast(f"📁 Документ визначено як: {mode}")
                    
                    # Гарантоване отримання промпта (виправляє KeyError)
                    prompt = prompts.get(mode, prompts["Загальна консультація"])
                    analysis = analyze_long_text(content, prompt)
                    
                    s_query = ""
                    if "Суд" in mode:
                        resp = client.chat.completions.create(
                            messages=[{"role": "user", "content": f"Напиши 3 ключові слова для пошуку практики (без лапок): {analysis[:300]}"}],
                            model="llama-3.1-8b-instant"
                        )
                        s_query = resp.choices[0].message.content.strip().replace('"', '')

                    name = url_input if "Посилання" in source else uploaded_file.name
                    save_to_history(name, analysis, content, s_query)
                    
                    st.session_state.analysis_result = analysis
                    st.session_state.full_content = content
                    st.session_state.search_query = s_query
                    st.success("Аналіз успішно завершено! Перейдіть у вкладку 'Висновки'.")
                    st.rerun()
            except Exception as e:
                st.error(f"Виникла помилка під час аналізу: {e}")

with tabs[1]:
    if st.session_state.analysis_result:
        st.markdown(f'<div class="result-card">{st.session_state.analysis_result}</div>', unsafe_allow_html=True)
        
        st.markdown("### 🛠️ Експорт та ресурси")
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("📥 Завантажити звіт (.docx)", data=create_docx(st.session_state.analysis_result), file_name="Legal_Mind_Report.docx", use_container_width=True)
        with c2:
            if st.session_state.search_query:
                q = st.session_state.search_query.replace(' ', '+')
                st.link_button("🔍 Схожа судова практика", f"https://www.google.com/search?q=site:reyestr.court.gov.ua+{q}", use_container_width=True)

        st.markdown("---")
        st.subheader("💬 Уточнити деталі у AI-адвоката")
        with st.form("chat_legal"):
            user_msg = st.text_input("Напишіть ваше запитання щодо цього документа:")
            if st.form_submit_button("Отримати відповідь") and user_msg:
                with st.spinner('Опрацювання запиту...'):
                    ans = client.chat.completions.create(
                        messages=[
                            {"role": "system", "content": "Ти досвідчений адвокат. Відповідай на основі наданого документа."},
                            {"role": "user", "content": f"Документ: {st.session_state.full_content[:10000]}\nЗапитання: {user_msg}"}
                        ],
                        model="llama-3.1-8b-instant"
                    )
                    st.info(f"⚖️ **Відповідь адвоката:**\n\n{ans.choices[0].message.content}")
    else:
        st.info("Результати аналізу з'являться тут після обробки документа у першій вкладці.")